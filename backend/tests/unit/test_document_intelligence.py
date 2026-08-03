"""Stage 1 — document intelligence (MS-2 gate).

Parsing accuracy and structure preservation are 15% of the assignment's grade,
and "the text all came out somewhere" does not satisfy it. These tests assert
structure at cell and hierarchy level, because a substring assertion passes on a
flattened table and that is exactly the failure worth catching.

Fixtures are real PDF/DOCX/PPTX files (``fixtures/documents/generate.py``), not
mocks — parser behaviour against synthetic input proves nothing.
"""

from __future__ import annotations

import asyncio
import functools
import hashlib
import inspect
import io
import os
import time
import zipfile
from pathlib import Path
from typing import Any
from uuid import uuid4

import httpx
import pytest

from api.deps import set_roster_builder, set_store
from api.main import create_app
from contracts.document import StructuredDocument
from core.storage.memory import InMemoryStore
from orchestration.pipeline import Roster
from stages.s1_document_intelligence import parsers
from stages.s1_document_intelligence import stage as stage_module
from stages.s1_document_intelligence.chunking import chunk_blocks, estimate_tokens
from stages.s1_document_intelligence.equations import is_probable_equation, to_latex
from stages.s1_document_intelligence.errors import (
    DocumentTooLarge,
    EmptyDocument,
    ParseFailure,
    ParseTimeout,
    TooManyPages,
    UnsupportedMediaType,
)
from stages.s1_document_intelligence.stage import parse_document
from stages.s1_document_intelligence.structure import TextSpan, infer_heading_level
from stages.stubs import STUB_STAGES

DOCS = Path(__file__).resolve().parents[1] / "fixtures" / "documents"

PDF = "application/pdf"
DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PPTX = "application/vnd.openxmlformats-officedocument.presentationml.presentation"
TXT = "text/plain"

LIMITS = {"max_bytes": 26_214_400, "max_pages": 300, "timeout_s": 30.0}


async def _parse(name: str, mime: str, **overrides: object) -> tuple[StructuredDocument, list]:
    return await parse_document(
        document_id=uuid4(),
        payload=(DOCS / name).read_bytes(),
        filename=name,
        mime=mime,
        **{**LIMITS, **overrides},  # type: ignore[arg-type]
    )


# ─────────────────────────────────────────────────────────── format coverage


@pytest.mark.parametrize(
    ("name", "mime"),
    [("physics.pdf", PDF), ("history.docx", DOCX), ("lesson.pptx", PPTX), ("notes.md", TXT)],
)
async def test_every_required_format_parses(name: str, mime: str) -> None:
    """FR-01: PDF, DOCX, PPT, and text are all named in the assignment."""
    document, chunks = await _parse(name, mime)
    assert document.blocks
    assert chunks
    assert document.metadata.word_count and document.metadata.word_count > 10


# ────────────────────────────────────────────────────── structure preservation


async def test_pdf_heading_hierarchy_is_recovered() -> None:
    """PDF has no heading markup; the hierarchy is inferred and must be right."""
    document, _ = await _parse("physics.pdf", PDF)

    assert len(document.outline) == 1
    root = document.outline[0]
    assert "Newton" in root.title
    assert root.level == 1
    assert [c.level for c in root.children] == [2, 2]
    assert [c.title for c in root.children] == ["5.1 The First Law", "5.2 The Second Law"]


async def test_table_headers_do_not_leak_into_the_outline() -> None:
    """A bold table header row satisfies every heading cue.

    Left unhandled it is promoted into the hierarchy and corrupts the section
    path of every block after it.
    """
    document, _ = await _parse("physics.pdf", PDF)
    titles = [node.title for node in document.outline[0].children]
    assert not any("Quantity" in t for t in titles), f"table header leaked: {titles}"


@pytest.mark.parametrize(
    ("name", "mime", "headers", "first_row"),
    [
        ("physics.pdf", PDF, ["Quantity", "Symbol", "Unit"], ["Force", "F", "N"]),
        ("history.docx", DOCX, ["Year", "Event"], ["1905", "Partition announced"]),
        ("lesson.pptx", PPTX, ["Stage", "Location"], ["Light reactions", "Thylakoid"]),
    ],
)
async def test_tables_survive_as_cells_not_prose(
    name: str, mime: str, headers: list[str], first_row: list[str]
) -> None:
    """Flattening is the silent failure of 'structure preservation'.

    Asserted at cell level: a substring check would pass on a flattened table.
    """
    document, _ = await _parse(name, mime)
    tables = [b for b in document.blocks if b.type == "table" and b.table]
    assert tables, f"no table recovered from {name}"

    table = tables[0].table
    assert table is not None
    assert table.headers == headers
    assert table.rows[0] == first_row
    assert all(len(row) == len(table.headers) for row in table.rows)


async def test_equations_are_detected_and_kept_renderable() -> None:
    """A physics chapter that loses its equations cannot produce a lesson plan."""
    document, _ = await _parse("physics.pdf", PDF)
    equations = [b for b in document.blocks if b.type == "equation"]
    assert equations
    assert any("=" in b.text for b in equations)
    # Raw text is always retained, so a failed conversion degrades rather than loses.
    assert all(b.text for b in equations)


async def test_section_paths_locate_every_block_in_the_document() -> None:
    """Section paths are what let an evidence span tell a teacher where to look."""
    document, _ = await _parse("physics.pdf", PDF)
    equation = next(b for b in document.blocks if b.type == "equation")
    assert equation.section_path[-1] == "5.2 The Second Law"
    assert "Newton" in equation.section_path[0]


async def test_docx_styles_drive_headings() -> None:
    document, _ = await _parse("history.docx", DOCX)
    headings = [(b.level, b.text) for b in document.blocks if b.type == "heading"]
    assert (1, "The Partition of Bengal") in headings
    assert (2, "Background") in headings


async def test_pptx_slide_titles_become_headings_and_notes_are_kept() -> None:
    document, _ = await _parse("lesson.pptx", PPTX)
    headings = [b.text for b in document.blocks if b.type == "heading"]
    assert "Photosynthesis" in headings
    assert any(b.page == 2 for b in document.blocks)


async def test_markdown_hashes_are_honoured_as_explicit_levels() -> None:
    document, _ = await _parse("notes.md", TXT)
    headings = [(b.level, b.text) for b in document.blocks if b.type == "heading"]
    assert (1, "Cell Biology") in headings
    assert (2, "Transport") in headings


# ──────────────────────────────────────────────────────────────── versatility


async def test_a_humanities_document_yields_no_equations_and_that_is_fine() -> None:
    """NFR-01. Absent content is correct content, not a parsing failure."""
    document, _ = await _parse("history.docx", DOCX)
    assert document.stats.equations == 0
    assert document.stats.paragraphs > 0
    assert document.stats.tables == 1


# ────────────────────────────────────────────────────────────────── chunking


async def test_chunks_never_split_a_table() -> None:
    """A table row separated from its headers is no longer a table."""
    document, chunks = await _parse("physics.pdf", PDF)
    table_block = next(b for b in document.blocks if b.type == "table")
    owners = [c for c in chunks if table_block.block_id in c.block_ids]
    assert len(owners) == 1, "table block spans multiple chunks"


async def test_chunks_carry_the_metadata_a_citation_needs() -> None:
    _, chunks = await _parse("physics.pdf", PDF)
    assert all(c.chunk_id for c in chunks)
    assert all(c.token_count > 0 for c in chunks)
    assert any(c.section_path for c in chunks)


async def test_chunk_ordinals_are_dense_and_ordered() -> None:
    _, chunks = await _parse("history.docx", DOCX)
    assert [c.ordinal for c in chunks] == list(range(len(chunks)))


def test_oversized_atomic_block_becomes_its_own_chunk_rather_than_being_cut() -> None:
    from contracts.document import Block, TableData

    huge = Block(
        block_id="b_0000",
        type="table",
        text="x " * 4000,
        char_start=0,
        char_end=8000,
        table=TableData(headers=["a"], rows=[["b"]]),
    )
    chunks = chunk_blocks(uuid4(), [huge])
    assert len(chunks) == 1
    assert estimate_tokens(chunks[0].text) > 800


# ───────────────────────────────────────────────────────────── determinism


async def test_parsing_is_deterministic() -> None:
    """Same bytes in, same structure out. Stage 1 makes no model calls."""
    first, chunks_a = await _parse("physics.pdf", PDF)
    second, chunks_b = await _parse("physics.pdf", PDF)

    strip = {"document_id"}
    assert first.model_dump(exclude=strip) == second.model_dump(exclude=strip)
    assert [c.text for c in chunks_a] == [c.text for c in chunks_b]


# ────────────────────────────────────────────────────────── safety & limits


async def test_oversized_document_is_rejected_before_parsing() -> None:
    with pytest.raises(DocumentTooLarge):
        await _parse("physics.pdf", PDF, max_bytes=100)


async def test_page_cap_is_enforced() -> None:
    with pytest.raises(TooManyPages):
        await _parse("physics.pdf", PDF, max_pages=0)


async def test_unsupported_media_type_is_rejected() -> None:
    with pytest.raises(UnsupportedMediaType):
        await _parse("notes.md", "application/x-executable")


async def test_malformed_pdf_is_rejected_cleanly_not_hung() -> None:
    """Correct magic bytes, unusable body — must be a rejection, never a 500."""
    with pytest.raises(ParseFailure):
        await _parse("malformed.pdf", PDF)


async def test_document_with_no_extractable_text_is_rejected() -> None:
    """A scanned PDF lands here. Honest failure beats a package built on nothing."""
    with pytest.raises(EmptyDocument):
        await _parse("empty.txt", TXT)


async def test_parse_timeout_is_bounded() -> None:
    """Converts an indefinite hang on hostile input into a bounded rejection."""
    with pytest.raises(ParseTimeout):
        await _parse("physics.pdf", PDF, timeout_s=0.0001)


async def test_errors_carry_a_machine_readable_payload() -> None:
    with pytest.raises(DocumentTooLarge) as excinfo:
        await _parse("physics.pdf", PDF, max_bytes=100)
    payload = excinfo.value.to_payload()
    assert payload["code"] == "document_too_large"
    assert payload["details"]["limit_bytes"] == 100


# ───────────────────────────────────────────────────── prompt-injection input


async def test_instruction_shaped_text_is_extracted_as_ordinary_content() -> None:
    """H-13. Document text is data.

    Stage 1's obligation is to extract it faithfully and mark it as content; the
    delimiting that stops it acting as an instruction happens at the LLM boundary.
    Faithful extraction is what makes that possible.
    """
    document, _ = await _parse("adversarial.txt", TXT)
    text = " ".join(b.text for b in document.blocks)
    assert "IGNORE ALL PREVIOUS INSTRUCTIONS" in text
    assert all(b.type in {"heading", "paragraph", "list", "equation"} for b in document.blocks)
    assert "Photosynthesis" in text


# ──────────────────────────────────────────────────────────── unit helpers


@pytest.mark.parametrize(
    ("text", "expected"),
    [
        ("F = m * a", True),
        ("E = mc²", True),
        ("6CO2 + 6H2O = C6H12O6 + 6O2", True),
        ("The force is equal to mass times acceleration.", False),
        ("This is the introduction to the chapter.", False),
        ("", False),
    ],
)
def test_equation_detection_separates_maths_from_prose(text: str, expected: bool) -> None:
    """Prose containing an equals sign is still prose."""
    assert is_probable_equation(text) is expected


def test_latex_normalisation_converts_symbols() -> None:
    assert to_latex("α = β × 2") == r"\alpha = \beta \times 2"
    assert to_latex("Just some ordinary sentence text here.") is None


def test_latex_normalisation_handles_sub_and_superscripts() -> None:
    assert to_latex("E = mc²") == "E = mc^{2}"
    assert to_latex("H₂O = x") == "H_{2}O = x"


def test_existing_latex_passes_through_untouched() -> None:
    """A wrong reconstruction is worse than none, so markup is never rewritten."""
    assert to_latex(r"\vec{F} = m\vec{a}") == r"\vec{F} = m\vec{a}"
    assert to_latex(r"$E = mc^2$") == "E = mc^2"


@pytest.mark.parametrize(
    ("text", "expected"),
    [
        ("3.2 Newton's Laws", 2),
        ("Chapter 5 Motion", 1),
        ("1.2.3 Deep Section", 3),
        ("This is an ordinary sentence that runs on and on and should not match.", None),
    ],
)
def test_numbered_headings_are_recognised_without_font_information(
    text: str, expected: int | None
) -> None:
    """Explicit numbering wins outright — plain text has no typography to read."""
    assert infer_heading_level(TextSpan(text=text), body_size=None) == expected


def test_sentence_punctuation_disqualifies_a_heading() -> None:
    span = TextSpan(text="A short bold line.", font_size=20.0, bold=True)
    assert infer_heading_level(span, body_size=11.0) is None


# ═══════════════════════════════════════════════════ ingestion hardening ═══
#
# Everything below concerns a hostile file rather than a badly-made one. The
# distinction matters: a badly-made file has to fail cleanly, but a hostile file
# has to fail cleanly *and cheaply*, before it can spend the resources that were
# the point of sending it.

_REAL_DOCX = (DOCS / "history.docx").read_bytes()


def _repack(source: bytes, extra: dict[str, bytes]) -> bytes:
    """A real DOCX with extra members grafted on.

    Built from a genuine fixture rather than from scratch so that what is being
    tested is the added hostility, not an incidental defect in a hand-rolled
    container.
    """
    out = io.BytesIO()
    with (
        zipfile.ZipFile(io.BytesIO(source)) as src,
        zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as dst,
    ):
        for info in src.infolist():
            dst.writestr(info.filename, src.read(info.filename))
        for name, data in extra.items():
            dst.writestr(name, data)
    return out.getvalue()


def _swap_document_xml(source: bytes, document_xml: bytes) -> bytes:
    out = io.BytesIO()
    with (
        zipfile.ZipFile(io.BytesIO(source)) as src,
        zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as dst,
    ):
        for info in src.infolist():
            data = (
                document_xml
                if info.filename == "word/document.xml"
                else src.read(info.filename)
            )
            dst.writestr(info.filename, data)
    return out.getvalue()


# ───────────────────────────────────────────────────── decompression bombs


async def test_a_docx_declaring_gigabytes_is_rejected_before_a_parser_sees_it(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """The finding this guard exists for.

    DOCX is a ZIP. A 1 MB upload can declare a gigabyte of members, and every
    byte-count limit upstream passes it because the upload really is 1 MB — the
    cost lands when ``python-docx`` hands the whole member to lxml as one string.
    Measured before the fix: a 2.18 MB DOCX reached 2.9 GB of resident memory.

    The rejection has to happen *before* the parser is called, so the parser is
    replaced with one that records having been reached. It must not be.
    """
    reached = []

    def tripwire(*args: object, **kwargs: object) -> list:
        reached.append(True)
        return []

    monkeypatch.setitem(stage_module.PARSER_BY_MIME, DOCX, tripwire)

    bomb = _repack(_REAL_DOCX, {"word/bomb.xml": b"\0" * (1024 * 1024 * 1024)})
    assert len(bomb) < 2 * 1024 * 1024, "the upload itself is small — that is the attack"

    with pytest.raises(DocumentTooLarge) as excinfo:
        await parse_document(
            document_id=uuid4(),
            payload=bomb,
            filename="bomb.docx",
            mime=DOCX,
            **LIMITS,  # type: ignore[arg-type]
        )

    assert not reached, "the parser was handed the bomb"
    payload = excinfo.value.to_payload()
    assert payload["code"] == "document_too_large"
    assert excinfo.value.http_status == 413
    assert payload["details"]["reason"] == "archive_uncompressed_size"


async def test_one_hyper_compressible_member_is_rejected_on_ratio_alone() -> None:
    """A bomb does not have to be big in aggregate to be a bomb.

    150 MB sits under the 200 MB total, so the total rule alone would pass this.
    The per-member ratio is what catches a single hostile part hiding inside an
    otherwise ordinary archive.
    """
    bomb = _repack(_REAL_DOCX, {"word/bomb.xml": b"\0" * (150 * 1024 * 1024)})

    with pytest.raises(DocumentTooLarge) as excinfo:
        await parse_document(
            document_id=uuid4(),
            payload=bomb,
            filename="bomb.docx",
            mime=DOCX,
            **LIMITS,  # type: ignore[arg-type]
        )
    details = excinfo.value.details
    assert details["reason"] == "archive_compression_ratio"
    assert details["ratio"] > 200


async def test_an_archive_of_thousands_of_members_is_rejected_on_count() -> None:
    """Each part cheap, the total ruinous — the third independent rule."""
    bomb = _repack(
        _REAL_DOCX, {f"word/part{i}.xml": b"<a/>" for i in range(2100)}
    )

    with pytest.raises(DocumentTooLarge) as excinfo:
        await parse_document(
            document_id=uuid4(),
            payload=bomb,
            filename="bomb.pptx",
            mime=PPTX,
            **LIMITS,  # type: ignore[arg-type]
        )
    assert excinfo.value.details["reason"] == "archive_member_count"


async def test_a_legitimate_docx_passes_the_archive_inspection_untouched() -> None:
    """The guard is worthless if it rejects real documents."""
    document, chunks = await _parse("history.docx", DOCX)
    assert document.blocks
    assert chunks


async def test_a_container_that_is_not_a_readable_archive_is_a_clean_rejection() -> None:
    """ZIP magic bytes, unusable directory — a rejection, never a 500."""
    with pytest.raises(ParseFailure):
        await parse_document(
            document_id=uuid4(),
            payload=b"PK\x03\x04" + b"\xff" * 400,
            filename="broken.docx",
            mime=DOCX,
            **LIMITS,  # type: ignore[arg-type]
        )


def test_archive_inspection_reads_the_directory_and_never_decompresses() -> None:
    """Cost is what makes this check usable at the front door.

    A gigabyte-declaring archive has to be judged in the time it takes to read a
    central directory, or the check becomes the denial of service it prevents.
    """
    bomb = _repack(_REAL_DOCX, {"word/bomb.xml": b"\0" * (1024 * 1024 * 1024)})

    start = time.monotonic()
    with pytest.raises(DocumentTooLarge):
        parsers.inspect_ooxml_archive(bomb)
    assert time.monotonic() - start < 0.5


# ──────────────────────────────────────────── a timeout that ends the work


def note_pid(pid_path: str) -> list:
    """Warm-up task: proves a child can start and run this module's code."""
    with open(pid_path, "w") as handle:
        handle.write(str(os.getpid()))
    return []


def spin_until_killed(pid_path: str) -> list:
    """Stands in for a parser driven into pathological work.

    Defined at module level because the child process has to import it by name.
    Self-limiting at 120s so that a regression which sends this down the
    in-thread fallback fails the suite instead of hanging it — a thread cannot be
    killed, which is the entire point being tested.
    """
    import time as _time

    with open(pid_path, "w") as handle:
        handle.write(str(os.getpid()))
    deadline = _time.monotonic() + 120
    while _time.monotonic() < deadline:
        pass
    return []


def _process_state(pid: int) -> str:
    """POSIX process state, or ``gone``. A killed-but-unreaped child reads Z."""
    try:
        stat = Path(f"/proc/{pid}/stat").read_text()
    except OSError:
        return "gone"
    try:
        return stat.split(") ", 1)[1].split()[0]
    except IndexError:  # pragma: no cover
        return "gone"


async def test_a_parse_timeout_ends_the_work_and_not_merely_the_wait(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """The finding: ``wait_for`` around ``to_thread`` cancels the awaiter only.

    Measured before the fix: a parse ran 67 seconds past a 3 second budget, still
    allocating, with nothing left waiting for it. A thread cannot be cancelled; a
    process can be killed, so the parse runs in one.

    Asserted on the child's own PID rather than on elapsed time, because
    returning promptly is exactly the symptom the broken version also had.
    """
    if stage_module._get_mp_context() is None:  # pragma: no cover - env dependent
        pytest.skip("no subprocess parse path available in this environment")

    limits = stage_module.ParseLimits()

    # A child re-imports this module before it can run anything in it. Time that
    # once, so the budget below is a budget on the *parse* rather than on module
    # import — otherwise this test measures the wrong thing on a slow machine.
    warm_file = tmp_path / "warm.pid"
    started = time.monotonic()
    await stage_module._run_parser(
        functools.partial(note_pid, str(warm_file)), timeout_s=120.0, limits=limits
    )
    startup_s = time.monotonic() - started
    assert warm_file.read_text(), "a child never ran"

    killed: list[int] = []
    original_terminate = stage_module._terminate

    def recording_terminate(executor: object) -> None:
        killed.extend(getattr(executor, "_processes", None) or {})
        original_terminate(executor)  # type: ignore[arg-type]

    monkeypatch.setattr(stage_module, "_terminate", recording_terminate)

    pid_file = tmp_path / "child.pid"
    budget = startup_s + 3.0
    started = time.monotonic()
    with pytest.raises(ParseTimeout) as excinfo:
        await stage_module._run_parser(
            functools.partial(spin_until_killed, str(pid_file)),
            timeout_s=budget,
            limits=limits,
        )
    elapsed = time.monotonic() - started

    assert excinfo.value.details["work_terminated"] is True, "fell back to the thread path"
    assert elapsed < budget + 10, f"timeout was not honoured promptly: {elapsed:.1f}s"
    assert killed, "no child process was killed"

    # The child that was running the parse, and any the executor still held, are
    # gone. A killed-but-unreaped child reads as a zombie, which is dead enough:
    # it holds no memory and burns no CPU.
    if pid_file.exists():
        killed.append(int(pid_file.read_text()))
    for pid in set(killed):
        for _ in range(100):
            if _process_state(pid) in {"gone", "Z", "X"}:
                break
            await asyncio.sleep(0.05)
        assert _process_state(pid) in {"gone", "Z", "X"}, (
            f"the parser survived its own timeout (pid {pid} is {_process_state(pid)})"
        )


async def test_parsing_still_works_when_no_process_pool_can_be_created(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """A deployment that forbids subprocesses must still start and still parse.

    The guarantee degrades honestly — the timeout then bounds the wait rather
    than the work, which is logged once — but the service does not fail to boot,
    and that is the difference between a hardening measure and an outage.
    """
    monkeypatch.setattr(stage_module, "_get_mp_context", lambda: None)

    document, chunks = await _parse("physics.pdf", PDF)
    assert document.blocks
    assert chunks


async def test_the_thread_fallback_still_bounds_the_request() -> None:
    """Even degraded, the caller gets a bounded rejection rather than a hang."""
    with pytest.raises(ParseTimeout) as excinfo:
        await stage_module._run_in_thread(
            functools.partial(
                parsers.parse_pdf, (DOCS / "physics.pdf").read_bytes(), max_pages=300
            ),
            timeout_s=0.0001,
        )
    assert excinfo.value.details["work_terminated"] is False


def test_the_child_initializer_never_raises_on_a_hostile_platform() -> None:
    """Limits are best-effort by design.

    A sandbox that refuses ``setrlimit`` must not turn every parse into a 500;
    the wall-clock kill is the control that always holds.
    """
    stage_module._apply_child_limits(2 ** 62, 10 ** 6)


def test_parse_limits_fall_back_to_the_module_defaults() -> None:
    """Settings that cannot be read must not decide the envelope by accident."""
    resolved = stage_module._limits_from_settings()
    assert resolved.archive_uncompressed_bytes >= parsers.DEFAULT_MAX_ARCHIVE_BYTES // 4
    assert resolved.max_blocks > 0 and resolved.max_chars > 0


# ────────────────────────────────────────── one choke point for all formats


@pytest.mark.parametrize(
    ("name", "mime"),
    [("physics.pdf", PDF), ("history.docx", DOCX), ("lesson.pptx", PPTX), ("notes.md", TXT)],
)
async def test_every_parser_is_capped_by_the_same_block_ceiling(
    name: str, mime: str
) -> None:
    """Four parsers, one ceiling.

    The bug this replaces: ``parse_docx(payload, **_)`` accepted ``max_pages``
    and dropped it on the floor, so DOCX had no volume limit at all while the
    caller believed it had one. Capping in ``_BlockBuilder.add`` — the single
    place every block of every format is created — is what makes "all four are
    capped" checkable rather than hopeful.
    """
    with pytest.raises(DocumentTooLarge) as excinfo:
        await _parse(name, mime, limits=stage_module.ParseLimits(max_blocks=2))
    assert excinfo.value.details["reason"] == "block_ceiling"
    assert excinfo.value.http_status == 413


@pytest.mark.parametrize(
    ("name", "mime"),
    [("physics.pdf", PDF), ("history.docx", DOCX), ("lesson.pptx", PPTX), ("notes.md", TXT)],
)
async def test_every_parser_is_capped_by_the_same_character_ceiling(
    name: str, mime: str
) -> None:
    """Block count alone does not bound memory; one enormous block would pass."""
    with pytest.raises(DocumentTooLarge) as excinfo:
        await _parse(name, mime, limits=stage_module.ParseLimits(max_chars=40))
    assert excinfo.value.details["reason"] == "character_ceiling"


def test_no_parser_can_silently_swallow_a_limit_it_was_handed() -> None:
    """The shape of the fix, asserted directly.

    ``**kwargs`` on a parser means a limit can be passed, accepted and ignored
    with nothing anywhere reporting it. None of the four may have one, and all
    four must name every limit they are given.
    """
    required = {"max_pages", "max_blocks", "max_chars"}
    for parser in (parsers.parse_pdf, parsers.parse_docx, parsers.parse_pptx, parsers.parse_text):
        signature = inspect.signature(parser)
        kinds = {p.kind for p in signature.parameters.values()}
        assert inspect.Parameter.VAR_KEYWORD not in kinds, f"{parser.__name__} swallows kwargs"
        assert required <= set(signature.parameters), f"{parser.__name__} is missing a limit"


async def test_the_real_ncert_chapter_stays_far_below_every_ceiling() -> None:
    """FAQ.md names an NCERT chapter as the benchmark input.

    A limit tuned so tightly that the document the product exists to read gets
    rejected is a broken limit, not a safe one. This pins the headroom so a later
    tightening cannot quietly cross it.
    """
    book = Path("/home/vikas/EduForge-AI/Books/leph101.pdf")
    if not book.exists():  # pragma: no cover - the corpus is not in git
        pytest.skip("NCERT benchmark chapter not present")

    document, chunks = await parse_document(
        document_id=uuid4(),
        payload=book.read_bytes(),
        filename="leph101.pdf",
        mime=PDF,
        max_bytes=26_214_400,
        max_pages=300,
        timeout_s=180.0,
    )
    characters = sum(len(b.text) for b in document.blocks)

    assert document.metadata.page_count == 44
    assert document.stats.tables > 0 and document.stats.equations > 0
    assert chunks
    # Two orders of magnitude of headroom on both ceilings.
    assert len(document.blocks) < parsers.DEFAULT_MAX_BLOCKS // 50
    assert characters < parsers.DEFAULT_MAX_TEXT_CHARS // 50


# ────────────────────────────────────────────────── DOCX honours its limit


def _docx_with_page_breaks(count: int) -> bytes:
    from docx import Document  # type: ignore[import-untyped]
    from docx.enum.text import WD_BREAK  # type: ignore[import-untyped]

    document = Document()
    document.add_heading("Chapter", level=1)
    for index in range(count):
        paragraph = document.add_paragraph(f"Body text for page {index}.")
        paragraph.add_run().add_break(WD_BREAK.PAGE)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


async def test_docx_now_honours_the_page_limit_it_used_to_discard() -> None:
    with pytest.raises(TooManyPages) as excinfo:
        await parse_document(
            document_id=uuid4(),
            payload=_docx_with_page_breaks(12),
            filename="long.docx",
            mime=DOCX,
            max_bytes=26_214_400,
            max_pages=5,
            timeout_s=30.0,
        )
    assert excinfo.value.details["page_count"] == 13


async def test_a_docx_that_declares_no_pages_is_not_rejected_on_a_guess() -> None:
    """A limit that fires on an estimate is worse than no limit.

    A DOCX has no page count until it is laid out. Where the file records none,
    the page rule stands down and the block and character ceilings do the work.
    """
    document, _ = await _parse("history.docx", DOCX, max_pages=1)
    assert document.blocks


# ─────────────────────────────────────────── XXE and entity expansion (H-15)
#
# OOXML parts are XML, so both classic XML attacks arrive dressed as documents.
# ``python-docx`` and ``python-pptx`` build their lxml parser with
# ``resolve_entities=False``, which is what defeats both — an inherited default
# rather than a decision this codebase made, and therefore exactly the kind of
# protection that disappears silently on a dependency bump. These tests assert
# the *behaviour*, so the day it changes the suite says so.

_XXE_DOCUMENT = b"""<?xml version="1.0"?>
<!DOCTYPE w:document [
  <!ENTITY xxe SYSTEM "file:///etc/passwd">
]>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:body><w:p><w:r><w:t>LEAKED:&xxe;</w:t></w:r></w:p></w:body>
</w:document>"""

_BILLION_LAUGHS = b"""<?xml version="1.0"?>
<!DOCTYPE w:document [
  <!ENTITY a "aaaaaaaaaa">
  <!ENTITY b "&a;&a;&a;&a;&a;&a;&a;&a;&a;&a;">
  <!ENTITY c "&b;&b;&b;&b;&b;&b;&b;&b;&b;&b;">
  <!ENTITY d "&c;&c;&c;&c;&c;&c;&c;&c;&c;&c;">
  <!ENTITY e "&d;&d;&d;&d;&d;&d;&d;&d;&d;&d;">
  <!ENTITY f "&e;&e;&e;&e;&e;&e;&e;&e;&e;&e;">
  <!ENTITY g "&f;&f;&f;&f;&f;&f;&f;&f;&f;&f;">
  <!ENTITY h "&g;&g;&g;&g;&g;&g;&g;&g;&g;&g;">
]>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:body><w:p><w:r><w:t>BOOM:&h;</w:t></w:r></w:p></w:body>
</w:document>"""


async def test_an_external_entity_in_a_docx_never_reaches_the_filesystem() -> None:
    """A document that asks to be told the contents of /etc/passwd is told nothing.

    If this ever fails, an uploaded file can read server-side files and the text
    it stole is handed to an LLM and then printed in a teacher's lesson plan.
    """
    payload = _swap_document_xml(_REAL_DOCX, _XXE_DOCUMENT)

    try:
        blocks = await stage_module._run_in_thread(
            functools.partial(parsers.parse_docx, payload, max_pages=300), timeout_s=30.0
        )
    except ParseFailure:
        return  # refusing to parse it at all is also a correct answer

    text = " ".join(b.text for b in blocks)
    assert "root:" not in text, "an external entity was resolved"
    assert "/bin/" not in text
    assert "LEAKED:" in text, "the surrounding text should still be extracted"


async def test_a_billion_laughs_docx_is_bounded_rather_than_expanded() -> None:
    """Ten nested entities, 10^8 characters if expanded. It must not be."""
    payload = _swap_document_xml(_REAL_DOCX, _BILLION_LAUGHS)

    started = time.monotonic()
    try:
        blocks = await stage_module._run_in_thread(
            functools.partial(parsers.parse_docx, payload, max_pages=300), timeout_s=30.0
        )
    except ParseFailure:
        assert time.monotonic() - started < 10
        return

    assert sum(len(b.text) for b in blocks) < 100_000, "entities were expanded"


async def test_the_same_defence_holds_for_pptx() -> None:
    """PPTX is the same container and the same parser family, so the same test."""
    from pptx import Presentation  # type: ignore[import-untyped]

    buffer = io.BytesIO()
    Presentation().save(buffer)
    deck = buffer.getvalue()

    hostile = b"""<?xml version="1.0"?>
<!DOCTYPE p:sld [ <!ENTITY xxe SYSTEM "file:///etc/passwd"> ]>
<p:sld xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"
       xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
<p:cSld><p:spTree>&xxe;</p:spTree></p:cSld></p:sld>"""

    out = io.BytesIO()
    with (
        zipfile.ZipFile(io.BytesIO(deck)) as src,
        zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as dst,
    ):
        for info in src.infolist():
            dst.writestr(info.filename, src.read(info.filename))
        dst.writestr("ppt/slides/slide99.xml", hostile)

    try:
        blocks = parsers.parse_pptx(out.getvalue(), max_pages=300)
    except (ParseFailure, TooManyPages):
        return
    assert "root:" not in " ".join(b.text for b in blocks)


# ──────────────────────────────────────────── the upload route's size check


async def _stub_roster(*_args: object) -> Roster:
    return Roster(stages=STUB_STAGES, llm=None)


@pytest.fixture
async def upload_client() -> Any:
    set_store(InMemoryStore())
    set_roster_builder(_stub_roster)
    transport = httpx.ASGITransport(app=create_app())
    async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
        yield client


async def test_a_declared_length_over_the_limit_is_rejected_on_the_header(
    upload_client: httpx.AsyncClient,
) -> None:
    """The cheapest rejection available.

    A client that announces four gigabytes is answered from its own header. The
    liar is still caught by the bounded read, so nothing is trusted here — this
    only avoids paying for the honest case.
    """
    boundary = "----eduforge"
    body = (
        f"--{boundary}\r\n"
        'Content-Disposition: form-data; name="file"; filename="a.pdf"\r\n'
        "Content-Type: application/pdf\r\n\r\n"
    ).encode() + b"%PDF-1.7\n" + b"x" * 64 + f"\r\n--{boundary}--\r\n".encode()

    response = await upload_client.post(
        "/api/v1/documents",
        content=body,
        headers={
            "content-type": f"multipart/form-data; boundary={boundary}",
            "content-length": str(4 * 1024 * 1024 * 1024),
        },
    )

    assert response.status_code == 413
    error = response.json()["error"]
    assert error["code"] == "document_too_large"
    assert error["details"]["size_bytes"] == 4 * 1024 * 1024 * 1024


async def test_an_oversized_body_is_rejected_even_when_the_header_lies(
    upload_client: httpx.AsyncClient,
) -> None:
    """The check that actually holds. A truthful header is not required."""
    oversized = b"%PDF-1.7\n" + b"x" * (26 * 1024 * 1024)
    response = await upload_client.post(
        "/api/v1/documents", files={"file": ("big.pdf", oversized, "application/pdf")}
    )
    assert response.status_code == 413
    assert response.json()["error"]["code"] == "document_too_large"


async def test_a_legitimate_upload_is_unaffected_by_the_bounded_read(
    upload_client: httpx.AsyncClient,
) -> None:
    """A file that spans several read chunks must be reassembled byte-exact."""
    content = b"%PDF-1.7\n" + os.urandom(700 * 1024)
    response = await upload_client.post(
        "/api/v1/documents", files={"file": ("real.pdf", content, "application/pdf")}
    )
    assert response.status_code == 201, response.text
    body = response.json()
    assert body["size_bytes"] == len(content)
    assert body["sha256"] == hashlib.sha256(content).hexdigest()


async def test_the_bounded_read_stops_instead_of_buffering_the_whole_body() -> None:
    """Unit-level proof that nothing past one chunk over the limit is held."""
    from api.routes.documents import _CHUNK_BYTES, _OverLimitError, _read_bounded

    class _Endless:
        """A body that never ends — the shape of the attack."""

        def __init__(self) -> None:
            self.served = 0

        async def read(self, size: int) -> bytes:
            self.served += size
            return b"\0" * size

    body = _Endless()
    with pytest.raises(_OverLimitError):
        await _read_bounded(body, 1024 * 1024)  # type: ignore[arg-type]
    assert body.served <= 1024 * 1024 + _CHUNK_BYTES
